# tests/test_core.py

def test_always_true():
    assert True
"""

import os
import tempfile
from pathlib import Path
from docx import Document
import pytest
from PIL import Image

from multivec.data.load.docx import DOCXLoader
from multivec.exceptions import DOCXLoaderError

@pytest.fixture
def sample_docx():
    # Create a temporary DOCX file for testing
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        doc = Document()
        doc.add_paragraph(
            "This is a test document with some keywords like Python, NLP, and Image Processing."
        )
        doc.add_picture(
            "tests/resources/test_image.png"
        )  # Ensure you have a test image in this path
        doc.save(tmp_file.name)
    yield tmp_file.name
    os.unlink(tmp_file.name)


@pytest.fixture
def processor(sample_docx):
    return DOCXLoader(sample_docx)


def test_initialization(sample_docx):
    processor = DOCXLoader(sample_docx)
    assert processor.docx_path == Path(sample_docx)
    assert processor.output_dir.exists()


def test_initialization_with_invalid_file():
    with pytest.raises(DOCXLoaderError):
        DOCXLoader("nonexistent.docx")


def test_initialization_with_wrong_file_type():
    with tempfile.NamedTemporaryFile(suffix=".txt") as tmp_file:
        with pytest.raises(DOCXLoaderError):
            DOCXLoader(tmp_file.name)


def test_extract_text(processor):
    texts = processor.extract_text()
    assert len(texts) > 0
    assert isinstance(texts[0].content, str)
    assert "test document" in texts[0].content


def test_extract_keywords(processor):
    text = "This is a test document about Python and NLP."
    keywords = processor.extract_keywords(text)
    assert len(keywords) > 0
    assert "Python" in keywords
    assert "NLP" in keywords


def test_highlight_keywords_in_image(processor, tmp_path):
    # Create a simple test image
    test_image_path = tmp_path / "test_image.png"
    Image.new("RGB", (100, 100), color="white").save(test_image_path)

    keywords = ["Test", "Keyword"]
    highlighted_path = processor.highlight_keywords_in_image(
        str(test_image_path), keywords
    )

    assert Path(highlighted_path).exists()
    assert "_highlighted" in highlighted_path


def test_process_images(processor):
    image_docs = processor.process_images()
    assert len(image_docs) > 0
    assert isinstance(image_docs[0].content, str)
    assert Path(image_docs[0].content).exists()
    assert "keywords" in image_docs[0].metadata


def test_process(processor):
    results = processor.process()
    assert "text" in results
    assert "images" in results
    assert len(results["text"]) > 0
    assert len(results["images"]) > 0


# Add more tests
"""