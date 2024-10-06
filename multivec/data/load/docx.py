import tempfile
from typing import List, Optional
from docx import Document
import pymupdf
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import spacy
from docx2pdf import convert
from spacy.util import get_package_path
from spacy.cli import download

from multivec.exceptions import DOCXLoaderError
from multivec.utils.base_format import ImageDocument, TextDocument

class DOCXLoader:
    """
    Load DOCX file, extract text and images, and highlight important keywords in images.
    """
    def __init__(self, docx_path: str, output_dir: Optional[str] = None, nlp_model: str = "en_core_web_sm"):
        self.docx_path = Path(docx_path)
        self.output_dir = Path(output_dir) if output_dir else Path(tempfile.mkdtemp())
        self._validate_input()
        self._ensure_nlp_model(nlp_model)
        self.docx = Document(self.docx_path)
        self.pdf_path = self._convert_to_pdf()
        self.nlp = spacy.load(nlp_model)

    def _ensure_nlp_model(self, model_name: str) -> None:
        """Ensure that the specified spaCy NLP model is available, and download if missing."""
        try:
            # Check if the model is installed by trying to get its package path
            get_package_path(model_name)
        except OSError:
            # If model is not installed, download it
            print(f"Downloading spaCy model '{model_name}'...")
            download(model_name)

    def _validate_input(self) -> None:
        """Validate input file and output directory."""
        if not self.docx_path.exists():
            raise DOCXLoaderError(f"File not found: {self.docx_path}")
        if not self.docx_path.suffix.lower() == '.docx':
            raise DOCXLoaderError(f"Invalid file format. Expected .docx, got {self.docx_path.suffix}")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _convert_to_pdf(self) -> Path:
        """Convert DOCX to PDF for image extraction."""
        pdf_path = self.output_dir / f"{self.docx_path.stem}.pdf"
        
        # Convert DOCX to PDF
        convert(self.docx_path, pdf_path)
        
        if not pdf_path.exists():
            raise DOCXLoaderError(f"Failed to convert DOCX to PDF: {pdf_path}")
        
        return pdf_path

    def extract_text(self) -> List[TextDocument]:
        """
        Extract text from the DOCX file.

        Returns:
            List[TextDocument]: A list of TextDocument objects.
        """
        texts = []
        for page_num, paragraph in enumerate(self.docx.paragraphs, start=1):
            if paragraph.text.strip():
                texts.append(
                    TextDocument(
                        content=paragraph.text,
                        metadata={"type": "text", "page": page_num},
                        page_index=page_num
                    )
                )
        return texts

    def extract_keywords(self, text: str, num_keywords: int = 5) -> List[str]:
        """
        Extract important keywords from the given text using spaCy.

        Args:
            text (str): The input text.
            num_keywords (int): The number of keywords to extract.

        Returns:
            List[str]: A list of extracted keywords.
        """
        doc = self.nlp(text)
        keywords = [token.text for token in doc if token.pos_ in ['NOUN', 'PROPN', 'ADJ']]
        return list(set(keywords))[:num_keywords]

    def highlight_keywords_in_image(self, image_path: str, keywords: List[str]) -> str:
        """
        Highlight keywords in the given image.

        Args:
            image_path (str): Path to the input image.
            keywords (List[str]): List of keywords to highlight.

        Returns:
            str: Path to the highlighted image.
        """
        with Image.open(image_path) as img:
            draw = ImageDraw.Draw(img)
            font = ImageFont.load_default()

            for i, keyword in enumerate(keywords):
                # Simple positioning, can be improved
                position = (10, 10 + i * 20)
                draw.text(position, keyword, fill="red", font=font)

            highlighted_path = f"{image_path.rsplit('.', 1)[0]}_highlighted.png"
            img.save(highlighted_path)

        return highlighted_path

    def process_images(self) -> List[ImageDocument]:
        """
        Extract images from the DOCX file, process them with NLP, and highlight keywords.

        Returns:
            List[ImageDocument]: A list of ImageDocument objects with highlighted keywords.
        """
        image_docs = []
        pdf_document = pymupdf.open(self.pdf_path)
        full_text = " ".join([para.text for para in self.docx.paragraphs])
        keywords = self.extract_keywords(full_text)

        for page_num, page in enumerate(pdf_document, start=1):
            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                image_path = self.output_dir / f"image_p{page_num}_i{img_index}.{base_image['ext']}"
                with open(image_path, "wb") as image_file:
                    image_file.write(image_bytes)

                highlighted_image_path = self.highlight_keywords_in_image(str(image_path), keywords)

                image_docs.append(
                    ImageDocument(
                        content=highlighted_image_path,
                        metadata={
                            "type": "image",
                            "page": page_num,
                            "format": base_image['ext'],
                            "size": len(image_bytes),
                            "keywords": keywords
                        },
                        page_index=page_num
                    )
                )

        pdf_document.close()
        return image_docs

    def process(self) -> dict:
        """
        Process the DOCX file, extract text and images, and highlight keywords in images.

        Returns:
            dict: A dictionary containing lists of TextDocument and ImageDocument objects.
        """
        return {
            "text": self.extract_text(),
            "images": self.process_images()
        }
